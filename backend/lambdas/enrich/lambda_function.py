"""
XO Platform - POST /enrich Lambda
Analyzes uploaded documents using Claude API.
Reads metadata from PostgreSQL, tracks enrichment in DB.

Two-phase async pattern:
  Phase 1 (synchronous): Auth, validate, create enrichment record, self-invoke async
  Phase 2 (async): Extract text, transcribe audio, analyze with Claude, write results
"""

import json
import os
import time
import re
import uuid
import boto3
import io
import csv
from datetime import datetime, timezone
from anthropic import Anthropic
from auth_helper import require_auth, get_db_connection, CORS_HEADERS
try:
    from crypto_helper import (
        decrypt, decrypt_json, unwrap_client_key,
        client_decrypt, client_decrypt_json,
        decrypt_s3_body, decrypt_s3_bytes, encrypt_s3_body
    )
except ImportError:
    import json as _json
    def decrypt(x): return x
    def decrypt_json(x):
        if not x: return None
        try: return _json.loads(x)
        except: return None
    def unwrap_client_key(x): return None
    def client_decrypt(k, x): return x
    def client_decrypt_json(k, x):
        if not x: return None
        try: return _json.loads(x)
        except: return None
    def decrypt_s3_body(k, b): return b if isinstance(b, str) else b.decode('utf-8', errors='replace') if b else ''
    def decrypt_s3_bytes(k, d): return d
    def encrypt_s3_body(k, b): return b if isinstance(b, bytes) else (b.encode('utf-8') if b else b'')

s3_client = boto3.client('s3')
lambda_client = boto3.client('lambda')
transcribe_client = boto3.client('transcribe')

BUCKET_NAME = os.environ.get('BUCKET_NAME', 'xo-client-data-mv')
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY')
FUNCTION_NAME = os.environ.get('AWS_LAMBDA_FUNCTION_NAME', 'xo-enrich')
STREAMLINE_WEBHOOK_URL = os.environ.get('STREAMLINE_WEBHOOK_URL', '')
AUDIO_EXTENSIONS = {'mp3', 'wav', 'm4a', 'aac', 'ogg', 'flac', 'wma', 'mp4', 'webm'}
SYSTEM_SKILLS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'system-skills')

anthropic_client = Anthropic(api_key=ANTHROPIC_API_KEY)


def _get_system_config_value(conn, key):
    """Read a single value from system_config table. Opens own connection if conn is None."""
    own_conn = False
    try:
        if conn is None:
            conn = get_db_connection()
            own_conn = True
        cur = conn.cursor()
        cur.execute("SELECT config_value FROM system_config WHERE config_key = %s", (key,))
        row = cur.fetchone()
        cur.close()
        if own_conn:
            conn.close()
        return row[0] if row else ''
    except Exception as e:
        print(f"Failed to read system_config key '{key}' (non-fatal): {e}")
        if own_conn and conn:
            try:
                conn.close()
            except Exception:
                pass
        return ''


def lambda_handler(event, context):
    """
    Two-phase enrichment handler.

    Phase 1 (API Gateway request): Auth, validate, create record, self-invoke async.
    Phase 2 (async invocation): Run full enrichment pipeline.
    """

    # Phase 2: async pipeline execution
    if event.get('_async_phase'):
        return _run_enrichment_pipeline(event)

    # Phase 1: synchronous API Gateway handler

    # Handle OPTIONS preflight
    if event.get('httpMethod') == 'OPTIONS':
        return {'statusCode': 200, 'headers': CORS_HEADERS, 'body': ''}

    # Route: POST /send-to-streamline
    resource = event.get('resource', '')
    if resource == '/send-to-streamline':
        return _handle_send_to_streamline(event)

    # Auth check
    user, err = require_auth(event)
    if err:
        return err

    conn = None

    try:
        body = json.loads(event.get('body', '{}'))
        client_id = body.get('client_id', '').strip()
        requested_model = body.get('model', '')

        if not client_id:
            return {
                'statusCode': 400,
                'headers': CORS_HEADERS,
                'body': json.dumps({'error': 'client_id is required'})
            }

        # Read client metadata from DB
        conn = get_db_connection()
        cur = conn.cursor()

        # Also read user's preferred_model as fallback
        cur.execute(
            "SELECT COALESCE(preferred_model, 'claude-sonnet-4-5-20250929') FROM users WHERE id = %s",
            (user['user_id'],)
        )
        user_row = cur.fetchone()
        db_model = user_row[0] if user_row else 'claude-sonnet-4-5-20250929'

        # Priority: request body > user preference > default
        allowed_models = ['claude-opus-4-6', 'claude-sonnet-4-5-20250929', 'claude-haiku-4-5-20251001']
        model_to_use = requested_model if requested_model in allowed_models else db_model
        print(f"Using model: {model_to_use}")

        cur.execute("""
            SELECT company_name, website_url, contact_name, contact_title,
                   contact_linkedin, industry, description, pain_point, id
            FROM clients
            WHERE s3_folder = %s AND user_id = %s
        """, (client_id, user['user_id']))

        row = cur.fetchone()
        if not row:
            cur.close()
            conn.close()
            return {
                'statusCode': 404,
                'headers': CORS_HEADERS,
                'body': json.dumps({'error': 'Client not found'})
            }

        db_client_id = str(row[8])

        # Query active upload S3 keys (only process active sources)
        cur.execute(
            "SELECT s3_key FROM uploads WHERE client_id = %s AND (status IS NULL OR status = 'active')",
            (db_client_id,)
        )
        active_keys = [r[0] for r in cur.fetchall()]
        print(f"Active upload keys for enrichment: {len(active_keys)}")

        # Create enrichment tracking record with stage
        cur.execute("""
            INSERT INTO enrichments (client_id, status, stage)
            VALUES (%s, 'processing', 'extracting')
            RETURNING id
        """, (db_client_id,))
        enrichment_id = str(cur.fetchone()[0])
        conn.commit()
        cur.close()
        conn.close()

        # Self-invoke async for the heavy lifting
        async_payload = {
            '_async_phase': True,
            'client_id': client_id,
            'db_client_id': db_client_id,
            'enrichment_id': enrichment_id,
            'user_id': user['user_id'],
            'model': model_to_use,
            'active_keys': active_keys
        }

        lambda_client.invoke(
            FunctionName=FUNCTION_NAME,
            InvocationType='Event',
            Payload=json.dumps(async_payload)
        )

        print(f"Async enrichment invoked for client: {client_id}, enrichment: {enrichment_id}")

        return {
            'statusCode': 200,
            'headers': CORS_HEADERS,
            'body': json.dumps({
                'job_id': client_id,
                'status': 'processing'
            })
        }

    except Exception as e:
        print(f"Error during enrichment setup: {str(e)}")
        import traceback
        traceback.print_exc()

        if conn:
            try:
                conn.close()
            except Exception:
                pass

        return {
            'statusCode': 500,
            'headers': CORS_HEADERS,
            'body': json.dumps({
                'error': 'Internal server error',
                'message': str(e)
            })
        }


def _run_enrichment_pipeline(event):
    """
    Phase 2: Async enrichment pipeline.
    Runs text extraction, audio transcription, Claude analysis, and writes results.
    """
    client_id = event['client_id']
    db_client_id = event['db_client_id']
    enrichment_id = event['enrichment_id']
    user_id = event['user_id']
    model = event.get('model', 'claude-sonnet-4-5-20250929')
    active_keys = event.get('active_keys', None)

    conn = None

    try:
        conn = get_db_connection()

        # Read client metadata
        cur = conn.cursor()
        cur.execute("""
            SELECT company_name, website_url, contact_name, contact_title,
                   contact_linkedin, industry, description, pain_point,
                   logo_s3_key, icon_s3_key,
                   COALESCE(streamline_webhook_enabled, FALSE),
                   contact_email, contact_phone, contacts_json, addresses_json,
                   streamline_webhook_url, encryption_key
            FROM clients WHERE id = %s
        """, (db_client_id,))
        row = cur.fetchone()
        if not row:
            update_enrichment_stage(conn, enrichment_id, 'error', status='error')
            conn.close()
            return {'status': 'error', 'message': 'Client not found'}

        # Unwrap per-client encryption key
        ck = unwrap_client_key(row[16]) if row[16] else None

        company_name = row[0] or 'Unknown Company'
        website = row[1] or ''
        contact_name = client_decrypt(ck, row[2]) or ''
        contact_title = client_decrypt(ck, row[3]) or ''
        contact_linkedin = client_decrypt(ck, row[4]) or ''
        industry = row[5] or ''
        description = row[6] or ''
        pain_point = row[7] or ''
        logo_s3_key = row[8] or ''
        icon_s3_key = row[9] or ''
        streamline_webhook_enabled = bool(row[10])
        contact_email = client_decrypt(ck, row[11]) or ''
        contact_phone = client_decrypt(ck, row[12]) or ''

        # Parse contacts_json with legacy fallback
        contacts_json_raw = row[13]
        contacts = []
        if contacts_json_raw:
            try:
                contacts = client_decrypt_json(ck, contacts_json_raw)
                if not contacts:
                    contacts = json.loads(contacts_json_raw)
            except (json.JSONDecodeError, TypeError):
                pass
        if not contacts:
            legacy = {'name': contact_name, 'title': contact_title, 'linkedin': contact_linkedin,
                      'email': contact_email, 'phone': contact_phone}
            if any(legacy.values()):
                contacts = [legacy]

        # Parse addresses_json
        addresses_json_raw = row[14]
        addresses = []
        if addresses_json_raw:
            try:
                addresses = client_decrypt_json(ck, addresses_json_raw)
                if not addresses:
                    addresses = json.loads(addresses_json_raw)
            except (json.JSONDecodeError, TypeError):
                pass

        streamline_webhook_url = client_decrypt(ck, row[15]) or ''
        if not streamline_webhook_url:
            streamline_webhook_url = _get_system_config_value(conn, 'enrichment_webhook_url')

        # Load system skills from DB first, fall back to bundled files
        system_skills = load_system_skills_from_db(cur)
        if not system_skills:
            print("No system skills in DB, falling back to bundled files")
            system_skills = load_system_skills()
        print(f"Loaded {len(system_skills)} system skills")

        # Read domain/client skills (decrypt with client key)
        skills = read_skills_from_db(cur, db_client_id, client_id, client_key=ck)
        print(f"Loaded {len(skills)} client skills for: {client_id}")
        cur.close()

        # Read client-config.md if it exists (decrypt with client key)
        client_config = read_client_config(client_id, client_key=ck)
        if client_config:
            print(f"Loaded client-config.md ({len(client_config)} chars)")
        else:
            print("No client-config.md found")

        # Stage: extracting (decrypt uploaded files with client key)
        update_enrichment_stage(conn, enrichment_id, 'extracting')
        extracted_text = extract_all_files(client_id, active_keys=active_keys, client_key=ck)
        print(f"Extracted text from {len(extracted_text)} files")

        # Stage: transcribing
        audio_files = find_audio_files(client_id)
        if audio_files:
            update_enrichment_stage(conn, enrichment_id, 'transcribing')
            print(f"Found {len(audio_files)} audio files to transcribe")
            transcripts = transcribe_audio_files(client_id, audio_files, client_key=ck)
            # Merge transcripts into extracted_text
            for filename, transcript in transcripts.items():
                extracted_text[filename] = transcript
            print(f"Transcribed {len(transcripts)} audio files")
        else:
            print("No audio files found, skipping transcription stage")

        if not extracted_text:
            update_enrichment_stage(conn, enrichment_id, 'error', status='error')
            conn.close()
            return {'status': 'error', 'message': 'No files found to analyze'}

        # Stage: researching (placeholder for future web research)
        update_enrichment_stage(conn, enrichment_id, 'researching')
        print("Research stage: placeholder (no web research yet)")

        # Stage: analyzing
        update_enrichment_stage(conn, enrichment_id, 'analyzing')
        print(f"Analyzing {len(extracted_text)} files for client: {client_id}")
        analysis = analyze_with_claude(
            company_name, website, contact_name, contact_title,
            contact_linkedin, industry, description, pain_point, extracted_text, skills,
            model=model, client_config=client_config, system_skills=system_skills,
            contacts=contacts
        )

        # Write results to S3 (encrypted with client key)
        results_key = f"{client_id}/results/analysis.json"
        results_body = json.dumps(analysis, indent=2)
        s3_client.put_object(
            Bucket=BUCKET_NAME,
            Key=results_key,
            Body=encrypt_s3_body(ck, results_body),
            ContentType='application/octet-stream'
        )

        # Stage: complete
        cur2 = conn.cursor()
        cur2.execute("""
            UPDATE enrichments
            SET status = 'complete', stage = 'complete', completed_at = NOW(), results_s3_key = %s
            WHERE id = %s
        """, (results_key, enrichment_id))
        conn.commit()
        cur2.close()
        conn.close()

        print(f"Analysis complete for client: {client_id}")

        # Fire webhook to Streamline (non-blocking — enrichment success is independent)
        # Per-client toggle takes priority; fall back to system-level toggle
        webhook_should_fire = streamline_webhook_enabled
        if not webhook_should_fire:
            sys_enabled = _get_system_config_value(None, 'streamline_webhook_enabled')
            webhook_should_fire = sys_enabled == 'true'
            if webhook_should_fire:
                print("Using system-level streamline_webhook_enabled=true (no per-client override)")

        if webhook_should_fire:
            _send_streamline_webhook(
                company_name=company_name,
                contacts=contacts,
                model=model,
                analysis=analysis,
                source_files=list(extracted_text.keys()),
                logo_s3_key=logo_s3_key,
                icon_s3_key=icon_s3_key,
                addresses=addresses,
                webhook_url=streamline_webhook_url
            )
        else:
            print("Streamline webhook disabled (per-client OFF, system default OFF)")

        return {'status': 'complete', 'client_id': client_id}

    except Exception as e:
        print(f"Error during enrichment pipeline: {str(e)}")
        import traceback
        traceback.print_exc()

        if conn:
            try:
                update_enrichment_stage(conn, enrichment_id, 'error', status='error')
            except Exception:
                pass
            try:
                conn.close()
            except Exception:
                pass

        return {'status': 'error', 'message': str(e)}


def _handle_send_to_streamline(event):
    """
    POST /send-to-streamline
    Re-sends the latest enrichment results to Streamline webhook
    without re-running enrichment.
    """
    user, err = require_auth(event)
    if err:
        return err

    conn = None
    try:
        body = json.loads(event.get('body', '{}'))
        client_id = body.get('client_id', '').strip()

        if not client_id:
            return {
                'statusCode': 400,
                'headers': CORS_HEADERS,
                'body': json.dumps({'error': 'client_id is required'})
            }

        # Read client metadata from DB
        conn = get_db_connection()
        cur = conn.cursor()

        cur.execute("""
            SELECT c.company_name, c.contact_name, c.contact_title, c.id,
                   e.results_s3_key, c.logo_s3_key, c.icon_s3_key,
                   c.contact_email, c.contact_phone, c.contacts_json,
                   c.contact_linkedin, c.addresses_json, c.streamline_webhook_url,
                   c.encryption_key
            FROM clients c
            LEFT JOIN enrichments e ON e.client_id = c.id AND e.status = 'complete'
            WHERE c.s3_folder = %s AND c.user_id = %s
            ORDER BY e.completed_at DESC
            LIMIT 1
        """, (client_id, user['user_id']))

        row = cur.fetchone()

        if not row:
            cur.close()
            conn.close()
            return {
                'statusCode': 404,
                'headers': CORS_HEADERS,
                'body': json.dumps({'error': 'Client or enrichment not found'})
            }

        # Unwrap per-client encryption key
        ck = unwrap_client_key(row[13]) if row[13] else None

        company_name = row[0] or 'Unknown Company'
        contact_name = client_decrypt(ck, row[1]) or ''
        contact_title = client_decrypt(ck, row[2]) or ''
        results_s3_key = row[4]
        logo_s3_key = row[5] or ''
        icon_s3_key = row[6] or ''
        contact_email = client_decrypt(ck, row[7]) or ''
        contact_phone = client_decrypt(ck, row[8]) or ''

        # Parse contacts_json with legacy fallback
        contacts_json_raw = row[9]
        contact_linkedin = client_decrypt(ck, row[10]) or ''
        contacts = []
        if contacts_json_raw:
            try:
                contacts = client_decrypt_json(ck, contacts_json_raw)
                if not contacts:
                    contacts = json.loads(contacts_json_raw)
            except (json.JSONDecodeError, TypeError):
                pass
        if not contacts:
            legacy = {'name': contact_name, 'title': contact_title, 'linkedin': contact_linkedin,
                      'email': contact_email, 'phone': contact_phone}
            if any(legacy.values()):
                contacts = [legacy]

        # Parse addresses_json
        addresses_json_raw = row[11]
        addresses = []
        if addresses_json_raw:
            try:
                addresses = client_decrypt_json(ck, addresses_json_raw)
                if not addresses:
                    addresses = json.loads(addresses_json_raw)
            except (json.JSONDecodeError, TypeError):
                pass

        manual_webhook_url = client_decrypt(ck, row[12]) or ''
        if not manual_webhook_url:
            manual_webhook_url = _get_system_config_value(conn, 'enrichment_webhook_url')

        cur.close()
        conn.close()

        if not results_s3_key:
            return {
                'statusCode': 404,
                'headers': CORS_HEADERS,
                'body': json.dumps({'error': 'No enrichment results found for this client'})
            }

        # Read results from S3
        s3_resp = s3_client.get_object(Bucket=BUCKET_NAME, Key=results_s3_key)
        analysis = json.loads(s3_resp['Body'].read().decode('utf-8'))

        source_files = analysis.get('analyzed_files', [])
        model_used = analysis.get('model', 'unknown')

        # Send to Streamline
        _send_streamline_webhook(
            company_name=company_name,
            contacts=contacts,
            model=model_used,
            analysis=analysis,
            source_files=source_files,
            logo_s3_key=logo_s3_key,
            icon_s3_key=icon_s3_key,
            addresses=addresses,
            webhook_url=manual_webhook_url
        )

        return {
            'statusCode': 200,
            'headers': CORS_HEADERS,
            'body': json.dumps({'message': 'Results sent to Streamline', 'status': 'sent'})
        }

    except Exception as e:
        print(f"Error in send-to-streamline: {str(e)}")
        if conn:
            try:
                conn.close()
            except Exception:
                pass
        return {
            'statusCode': 500,
            'headers': CORS_HEADERS,
            'body': json.dumps({'error': str(e)})
        }


def update_enrichment_stage(conn, enrichment_id, stage, status=None):
    """Update the enrichment stage (and optionally status) in DB."""
    try:
        cur = conn.cursor()
        if status:
            cur.execute("""
                UPDATE enrichments SET stage = %s, status = %s, completed_at = NOW()
                WHERE id = %s
            """, (stage, status, enrichment_id))
        else:
            cur.execute("""
                UPDATE enrichments SET stage = %s WHERE id = %s
            """, (stage, enrichment_id))
        conn.commit()
        cur.close()
        print(f"Stage updated: {stage}")
    except Exception as e:
        print(f"Error updating stage to {stage}: {e}")


def find_audio_files(client_id):
    """List audio files in the client's uploads folder."""
    audio_files = []
    try:
        response = s3_client.list_objects_v2(
            Bucket=BUCKET_NAME,
            Prefix=f"{client_id}/uploads/"
        )
        if 'Contents' not in response:
            return audio_files

        for obj in response['Contents']:
            key = obj['Key']
            filename = key.split('/')[-1]
            if not filename:
                continue
            ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
            if ext in AUDIO_EXTENSIONS:
                audio_files.append(key)
    except Exception as e:
        print(f"Error listing audio files: {e}")
    return audio_files


def read_audio_context(client_id, filename):
    """Read {filename}.context.json from S3 uploads folder if it exists."""
    context_key = f"{client_id}/uploads/{filename}.context.json"
    try:
        response = s3_client.get_object(Bucket=BUCKET_NAME, Key=context_key)
        return json.loads(response['Body'].read().decode('utf-8'))
    except Exception:
        return None


def transcribe_audio_files(client_id, audio_s3_keys, client_key=None):
    """Transcribe all audio files using AWS Transcribe. Returns {filename: transcript_text}."""
    transcripts = {}
    for s3_key in audio_s3_keys:
        filename = s3_key.split('/')[-1]
        try:
            transcript = transcribe_single_file(client_id, s3_key, filename)
            if transcript:
                # Build header with context if available
                ctx = read_audio_context(client_id, filename)
                if ctx:
                    header_parts = [f"Audio Transcript ({filename})"]
                    if ctx.get('date'):
                        header_parts.append(f"Date: {ctx['date']}")
                    if ctx.get('participants'):
                        header_parts.append(f"Participants: {ctx['participants']}")
                    if ctx.get('topic'):
                        header_parts.append(f"Topic: {ctx['topic']}")
                    header = " -- ".join(header_parts)
                else:
                    header = f"Audio Transcript ({filename})"

                transcripts[filename] = f"{header}\n\n{transcript}"
                # Save transcript to extracted folder (encrypted)
                transcript_key = f"{client_id}/extracted/{filename}.transcript.txt"
                s3_client.put_object(
                    Bucket=BUCKET_NAME,
                    Key=transcript_key,
                    Body=encrypt_s3_body(client_key, transcript),
                    ContentType='application/octet-stream'
                )
                print(f"Saved transcript: {transcript_key}")
        except Exception as e:
            print(f"Error transcribing {filename}: {e}")
    return transcripts


def transcribe_single_file(client_id, s3_key, filename):
    """Start and poll an AWS Transcribe job for a single audio file."""
    # Build job name: xo-{client_id_suffix}-{safe_filename}-{uuid}
    client_suffix = client_id[-12:] if len(client_id) > 12 else client_id
    safe_filename = re.sub(r'[^a-zA-Z0-9._-]', '_', filename.rsplit('.', 1)[0])
    job_name = f"xo-{client_suffix}-{safe_filename}-{uuid.uuid4().hex[:8]}"
    job_name = job_name[:200]  # Transcribe max job name length

    ext = filename.lower().rsplit('.', 1)[-1]
    media_format_map = {
        'mp3': 'mp3', 'wav': 'wav', 'm4a': 'mp4', 'aac': 'mp4',
        'ogg': 'ogg', 'flac': 'flac', 'wma': 'mp3',
        'mp4': 'mp4', 'webm': 'webm'
    }
    media_format = media_format_map.get(ext, 'mp3')

    media_uri = f"s3://{BUCKET_NAME}/{s3_key}"
    output_key = f"{client_id}/extracted/.transcribe-output/{filename}.json"

    print(f"Starting transcription job: {job_name} for {filename}")

    transcribe_client.start_transcription_job(
        TranscriptionJobName=job_name,
        Media={'MediaFileUri': media_uri},
        MediaFormat=media_format,
        LanguageCode='en-US',
        OutputBucketName=BUCKET_NAME,
        OutputKey=output_key
    )

    # Poll for completion (max 360s — video files can take 2-5 min)
    max_wait = 360
    elapsed = 0
    while elapsed < max_wait:
        time.sleep(5)
        elapsed += 5

        response = transcribe_client.get_transcription_job(
            TranscriptionJobName=job_name
        )
        status = response['TranscriptionJob']['TranscriptionJobStatus']

        if status == 'COMPLETED':
            print(f"Transcription complete: {job_name} ({elapsed}s)")
            return read_transcribe_output(output_key)
        elif status == 'FAILED':
            reason = response['TranscriptionJob'].get('FailureReason', 'Unknown')
            print(f"Transcription failed: {job_name} - {reason}")
            return None

        print(f"Transcription in progress: {job_name} ({elapsed}s)")

    print(f"Transcription timed out: {job_name} after {max_wait}s")
    return None


def read_transcribe_output(output_key):
    """Read and parse AWS Transcribe output JSON from S3."""
    try:
        response = s3_client.get_object(Bucket=BUCKET_NAME, Key=output_key)
        output = json.loads(response['Body'].read().decode('utf-8'))
        transcripts = output.get('results', {}).get('transcripts', [])
        if transcripts:
            return transcripts[0].get('transcript', '')
        return ''
    except Exception as e:
        print(f"Error reading transcribe output ({output_key}): {e}")
        return None


def load_system_skills_from_db(cur):
    """Load system skills from DB (client_id IS NULL), with S3 fallback for content."""
    skills = []
    try:
        cur.execute("SELECT name, content, s3_key FROM skills WHERE client_id IS NULL ORDER BY name")
        for row in cur.fetchall():
            name, content, s3_key = row[0], row[1], row[2]
            if not content and s3_key:
                try:
                    obj = s3_client.get_object(Bucket=BUCKET_NAME, Key=s3_key)
                    content = obj['Body'].read().decode('utf-8')
                except Exception as e:
                    print(f"Failed to load system skill from S3 ({s3_key}): {e}")
                    continue
            if content:
                skills.append({'name': name, 'content': content})
    except Exception as e:
        print(f"Error loading system skills from DB: {e}")
    return skills


def load_system_skills():
    """Load system skills from the bundled system-skills/ directory (fallback)."""
    skills = []
    try:
        if not os.path.isdir(SYSTEM_SKILLS_DIR):
            print(f"System skills directory not found: {SYSTEM_SKILLS_DIR}")
            return skills
        for filename in sorted(os.listdir(SYSTEM_SKILLS_DIR)):
            if not filename.endswith('.md'):
                continue
            filepath = os.path.join(SYSTEM_SKILLS_DIR, filename)
            with open(filepath, 'r') as f:
                content = f.read()
            skills.append({
                'name': filename.replace('.md', ''),
                'content': content
            })
    except Exception as e:
        print(f"Error loading system skills: {e}")
    return skills


def read_client_config(client_id, client_key=None):
    """Read client-config.md from S3 if it exists. Decrypts with client key."""
    try:
        response = s3_client.get_object(
            Bucket=BUCKET_NAME,
            Key=f"{client_id}/client-config.md"
        )
        raw = response['Body'].read()
        return decrypt_s3_body(client_key, raw)
    except Exception:
        return None


def read_skills_from_db(cur, db_client_id, client_id, client_key=None):
    """Read skills from DB, with S3 fallback for s3_key-only skills."""
    skills = []

    try:
        cur.execute("""
            SELECT name, content, s3_key FROM skills WHERE client_id = %s
        """, (str(db_client_id),))

        for row in cur.fetchall():
            name, content, s3_key = row

            if content:
                skills.append({'name': name, 'content': content})
            elif s3_key:
                # Fallback: read from S3 (decrypt with client key)
                try:
                    file_obj = s3_client.get_object(Bucket=BUCKET_NAME, Key=s3_key)
                    raw = file_obj['Body'].read()
                    skill_content = decrypt_s3_body(client_key, raw)
                    skills.append({'name': name, 'content': skill_content})
                except Exception as e:
                    print(f"Error reading skill from S3 ({s3_key}): {e}")

    except Exception as e:
        print(f"Error reading skills from DB: {e}")
        # Fallback to S3 listing
        skills = read_skills_from_s3(client_id, client_key=client_key)

    return skills


def read_skills_from_s3(client_id, client_key=None):
    """Fallback: Read all skill files from S3 skills folder."""
    skills = []

    try:
        response = s3_client.list_objects_v2(
            Bucket=BUCKET_NAME,
            Prefix=f"{client_id}/skills/"
        )

        if 'Contents' not in response:
            return skills

        for obj in response['Contents']:
            key = obj['Key']
            filename = key.split('/')[-1]

            if not filename or not filename.endswith('.md'):
                continue

            file_obj = s3_client.get_object(Bucket=BUCKET_NAME, Key=key)
            raw = file_obj['Body'].read()
            skill_content = decrypt_s3_body(client_key, raw)
            skills.append({
                'name': filename.replace('.md', ''),
                'content': skill_content
            })

    except Exception as e:
        print(f"Error reading skills from S3: {e}")

    return skills


def extract_all_files(client_id, active_keys=None, client_key=None):
    """Extract text from all uploaded files. If active_keys provided, only process those."""
    extracted_text = {}
    active_keys_set = set(active_keys) if active_keys else None

    try:
        response = s3_client.list_objects_v2(
            Bucket=BUCKET_NAME,
            Prefix=f"{client_id}/uploads/"
        )

        if 'Contents' not in response:
            return extracted_text

        for obj in response['Contents']:
            key = obj['Key']
            filename = key.split('/')[-1]

            if not filename or filename == '':
                continue

            # Skip context metadata files (uploaded alongside audio)
            if filename.endswith('.context.json'):
                continue

            # Filter by active keys if provided
            if active_keys_set is not None and key not in active_keys_set:
                print(f"Skipping inactive/deleted file: {filename}")
                continue

            # Skip audio files — handled by Transcribe stage
            ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
            if ext in AUDIO_EXTENSIONS:
                print(f"Audio file — will be handled by Transcribe: {filename}")
                continue

            print(f"Processing file: {filename}")

            file_obj = s3_client.get_object(Bucket=BUCKET_NAME, Key=key)
            file_content = file_obj['Body'].read()

            # Decrypt if encrypted with client key
            file_content = decrypt_s3_bytes(client_key, file_content)

            text = extract_text(filename, file_content)
            if text:
                extracted_text[filename] = text

    except Exception as e:
        print(f"Error extracting files: {str(e)}")

    return extracted_text


def extract_text(filename, file_content):
    """Extract text from different file types"""
    ext = filename.lower().split('.')[-1]

    try:
        if ext == 'csv':
            return extract_csv(file_content)
        elif ext == 'txt':
            return file_content.decode('utf-8')
        elif ext in ['xlsx', 'xls']:
            return extract_excel(file_content)
        elif ext == 'pdf':
            return extract_pdf(file_content)
        elif ext in ['docx', 'doc']:
            return extract_docx(file_content)
        else:
            print(f"Unsupported file type: {ext}")
            return None

    except Exception as e:
        print(f"Error extracting {filename}: {str(e)}")
        return None


def extract_csv(file_content):
    """Extract text from CSV file"""
    try:
        content = file_content.decode('utf-8')
        reader = csv.reader(io.StringIO(content))
        rows = list(reader)

        text = "CSV Data:\n"
        text += f"Total rows: {len(rows)}\n\n"

        if rows:
            text += "Header: " + ", ".join(rows[0]) + "\n\n"
            text += "Sample data (first 10 rows):\n"
            for i, row in enumerate(rows[1:11]):
                text += f"Row {i+1}: " + ", ".join(row) + "\n"

        return text

    except Exception as e:
        print(f"Error parsing CSV: {str(e)}")
        return str(file_content[:1000])


def extract_excel(file_content):
    """Extract text from Excel file"""
    try:
        import openpyxl
        from io import BytesIO

        wb = openpyxl.load_workbook(BytesIO(file_content), data_only=True)
        text = "Excel Data:\n\n"

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text += f"Sheet: {sheet_name}\n"
            text += f"Rows: {sheet.max_row}, Columns: {sheet.max_column}\n\n"

            text += "Sample data (first 10 rows):\n"
            for i, row in enumerate(sheet.iter_rows(max_row=10, values_only=True)):
                text += f"Row {i+1}: " + ", ".join([str(cell) if cell else "" for cell in row]) + "\n"
            text += "\n"

        return text

    except Exception as e:
        print(f"Error parsing Excel: {str(e)}")
        return f"Excel file with {len(file_content)} bytes"


def extract_pdf(file_content):
    """Extract text from PDF file"""
    try:
        from pypdf import PdfReader
        from io import BytesIO

        pdf = PdfReader(BytesIO(file_content))
        text = f"PDF Document ({len(pdf.pages)} pages):\n\n"

        for i, page in enumerate(pdf.pages[:10]):
            page_text = page.extract_text()
            if page_text:
                text += f"--- Page {i+1} ---\n{page_text}\n\n"

        return text

    except Exception as e:
        print(f"Error parsing PDF: {str(e)}")
        return f"PDF file with {len(file_content)} bytes"


def extract_docx(file_content):
    """Extract text from Word .docx file"""
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(file_content))
        text = f"Word Document ({len(doc.paragraphs)} paragraphs):\n\n"

        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Also extract text from tables
        if doc.tables:
            text += f"\n--- Tables ({len(doc.tables)}) ---\n"
            for i, table in enumerate(doc.tables):
                text += f"\nTable {i+1}:\n"
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    text += " | ".join(cells) + "\n"

        return text

    except Exception as e:
        print(f"Error parsing DOCX: {str(e)}")
        return f"Word document with {len(file_content)} bytes"


def _send_streamline_webhook(company_name, contacts, model, analysis, source_files, logo_s3_key='', icon_s3_key='', addresses=None, webhook_url=''):
    """
    POST enrichment results to Streamline webhook URL.
    Uses per-client webhook_url if provided, falls back to STREAMLINE_WEBHOOK_URL env var.
    Non-blocking — logs result but never fails the enrichment.
    """
    url = webhook_url or STREAMLINE_WEBHOOK_URL
    if not url:
        print("Streamline webhook: no URL configured (no per-client URL or env var), skipping")
        return

    try:
        import urllib.request

        primary = contacts[0] if contacts else {}
        contact_first = primary.get('firstName', '')
        contact_last = primary.get('lastName', '')
        # Fallback to legacy "name" field if firstName not present
        if not contact_first and primary.get('name'):
            parts = primary['name'].split(' ', 1)
            contact_first = parts[0]
            contact_last = parts[1] if len(parts) > 1 else ''
        contact_name = f"{contact_first} {contact_last}".strip()
        contact_title = primary.get('title', '')
        contact_display = contact_name
        if contact_title:
            contact_display += f" ({contact_title})"

        # Build full contacts array for payload
        contacts_payload = []
        for idx, c in enumerate(contacts):
            c_first = c.get('firstName', '')
            c_last = c.get('lastName', '')
            if not c_first and c.get('name'):
                parts = c['name'].split(' ', 1)
                c_first = parts[0]
                c_last = parts[1] if len(parts) > 1 else ''
            c_full = f"{c_first} {c_last}".strip()
            display = c_full
            if c.get('title'):
                display += f" ({c['title']})"
            contacts_payload.append({
                "first_name": c_first,
                "last_name": c_last,
                "name": c_full,
                "title": c.get('title', ''),
                "email": c.get('email', ''),
                "phone": c.get('phone', ''),
                "linkedin": c.get('linkedin', ''),
                "display": display
            })

        payload = {
            "client_name": company_name,
            # Legacy flat fields from primary contact
            "client_contact": contact_display,
            "client_contact_first_name": contact_first,
            "client_contact_last_name": contact_last,
            "client_email": primary.get('email', ''),
            "client_phone": primary.get('phone', ''),
            # Full contacts array
            "contacts": contacts_payload,
            # Addresses array
            "addresses": addresses or [],
            "enrichment_model": model,
            "enrichment_date": datetime.now(timezone.utc).isoformat(),
            "executive_summary": analysis.get("summary", ""),
            "problems_identified": analysis.get("problems", []),
            "proposed_schema": analysis.get("schema", {}).get("tables", []),
            "action_plan": {
                phase.get("phase", ""): phase.get("actions", [])
                for phase in analysis.get("plan", [])
            },
            "client_summary": analysis.get("client_summary", ""),
            "streamline_applications": analysis.get("streamline_applications", ""),
            "data_sources": analysis.get("sources", []),
            "source_files": source_files,
            "xo_results_url": "https://d2np82m8rfcd6u.cloudfront.net",
            "client_logo_url": f"https://xo-client-data-mv.s3.us-west-1.amazonaws.com/{logo_s3_key}" if logo_s3_key else None,
            "client_icon_url": f"https://xo-client-data-mv.s3.us-west-1.amazonaws.com/{icon_s3_key}" if icon_s3_key else None
        }

        data = json.dumps(payload).encode('utf-8')
        req = urllib.request.Request(
            url,
            data=data,
            headers={'Content-Type': 'application/json'},
            method='POST'
        )
        resp = urllib.request.urlopen(req, timeout=10)
        print(f"Streamline webhook: {resp.status} ({len(data)} bytes sent)")

    except Exception as e:
        print(f"Streamline webhook failed (non-fatal): {str(e)}")


def analyze_with_claude(company_name, website, contact_name, contact_title,
                        contact_linkedin, industry, description, pain_point, extracted_text, skills=None,
                        model='claude-sonnet-4-5-20250929', client_config=None, system_skills=None,
                        contacts=None):
    """
    Call Claude API with XO Capture Analysis prompt.
    Returns structured analysis JSON.
    """

    files_summary = "\n\n".join([
        f"=== {filename} ===\n{text[:5000]}"
        for filename, text in extracted_text.items()
    ])

    enrichment_info = []
    if website:
        enrichment_info.append(f"Company Website: {website}")

    # Build contact info from contacts array (preferred) or legacy fields
    if contacts and len(contacts) > 0:
        primary = contacts[0]
        display = primary.get('name', '')
        if primary.get('title'):
            display += f" ({primary['title']})"
        if display:
            enrichment_info.append(f"Primary Contact: {display}")
        if primary.get('linkedin'):
            enrichment_info.append(f"Primary Contact LinkedIn: {primary['linkedin']}")
        if primary.get('email'):
            enrichment_info.append(f"Primary Contact Email: {primary['email']}")
        if primary.get('phone'):
            enrichment_info.append(f"Primary Contact Phone: {primary['phone']}")
        for idx, c in enumerate(contacts[1:], start=2):
            display = c.get('name', '')
            if c.get('title'):
                display += f" ({c['title']})"
            if display:
                enrichment_info.append(f"Contact {idx}: {display}")
            if c.get('email'):
                enrichment_info.append(f"Contact {idx} Email: {c['email']}")
    else:
        if contact_name:
            enrichment_info.append(f"Primary Contact: {contact_name}" + (f" ({contact_title})" if contact_title else ""))
        if contact_linkedin:
            enrichment_info.append(f"Contact LinkedIn: {contact_linkedin}")

    if industry:
        enrichment_info.append(f"Industry: {industry}")
    if description:
        enrichment_info.append(f"Description: {description}")
    if pain_point:
        enrichment_info.append(f"Immediate Pain Point: {pain_point}")

    enrichment_section = "\n".join(enrichment_info) if enrichment_info else "Not provided"

    # System skills (always injected first -- analysis framework, output format, authority, enrichment process)
    system_skills_section = ""
    if system_skills and len(system_skills) > 0:
        system_skills_section = "\n\nSYSTEM INSTRUCTIONS (follow these rules for every analysis):\n"
        for skill in system_skills:
            system_skills_section += f"=== {skill['name'].upper().replace('-', ' ')} ===\n{skill['content']}\n\n"

    # Client config section (persistent context, like a CLAUDE.md for this client)
    config_section = ""
    if client_config:
        config_section = f"\n\nCLIENT CONFIGURATION (persistent context for this engagement):\n{client_config}\n"

    # Domain/client skills (editable per client, override defaults)
    skills_section = ""
    if skills and len(skills) > 0:
        skills_section = "\n\nCLIENT-SPECIFIC SKILLS & INSTRUCTIONS:\n"
        skills_section += "The following skills provide domain-specific knowledge and analysis instructions. Follow these instructions carefully -- they refine the system instructions above:\n\n"
        for skill in skills:
            skills_section += f"=== SKILL: {skill['name']} ===\n{skill['content']}\n\n"

    pain_point_section = ""
    if pain_point:
        pain_point_section = f"\n\nPRIORITY: The client has identified this as their immediate pain point: '{pain_point}'. Make this the #1 problem in your analysis. Lead the executive summary with it, ensure it appears first in the problems list with specific evidence and a concrete recommendation, and front-load the 7-day action plan with steps that directly address it."

    prompt = f"""You are an MBA-level business analyst conducting an XO Capture analysis. You have been given access to internal documents from a company. Analyze this business and provide strategic insights.
{system_skills_section}
COMPANY INFORMATION:
Company Name: {company_name}
{enrichment_section}
{config_section}
CLIENT DATA (Uploaded Documents):
{files_summary}
{skills_section}
TASK:
Analyze this business like an MBA analyst presenting on Monday morning.{pain_point_section}

Provide your analysis in structured, technical format. Follow these formatting rules strictly:

1. EXECUTIVE SUMMARY
   - Lead with the single biggest finding in the first sentence
   - 2-3 paragraphs maximum
   - Reference specific data points, not generalities

2. PROBLEMS IDENTIFIED (top 3-5)
   For each problem provide:
   - Title (clear, specific)
   - Severity (high/medium/low)
   - Evidence: cite specific data from the documents (row counts, dollar amounts, percentages)
   - Recommendation: concrete action with expected outcome

3. PROPOSED ARCHITECTURE
   - Provide an ASCII diagram showing the proposed system architecture using box-drawing characters (+, -, |, v, >)
   - Show data flow between components
   - Example format:
     +----------+     +----------+     +---------+
     | Source A  |---->| Process  |---->| Output  |
     +----------+     +----------+     +---------+

4. PROPOSED DATA SCHEMA
   - For each table, use this format:
     Table: table_name -- purpose
     | Column | Type | Description |
     |--------|------|-------------|
     | id | UUID | Primary key |
   - Show relationships between tables after the table definitions

5. 7/14/21 DAY ACTION PLAN
   - 7-day: Build and demo -- prototype the solution to the primary pain point, get it on screen, show it live
   - 14-day: Validate and connect -- incorporate feedback, validate data connections, prepare for real deployment
   - 21-day: Deploy or decide -- go live with the solution or make the build/buy decision
   - Numbered items within each phase
   - Each action should be specific and measurable
   - Include expected cost or effort level where possible

6. BOTTOM LINE
   - One paragraph: what to do first, what it will cost, what outcome to expect
   - Be direct -- this is the slide the CEO reads

7. CLIENT SUMMARY (XO Summary for Client)
   - A concise, client-ready summary that could be shared directly with the client
   - Open with: "Based on the information provided, XO has identified the following opportunities for [Company Name]:"
   - 3-5 bullet points, each a clear value proposition framed as a business outcome
   - Use plain business language -- no technical jargon, no architecture, no tools
   - NEVER include cost estimates, pricing, timelines, or build specifications
   - NEVER reference internal tools, frameworks, or technology stack
   - Frame everything as business outcomes and operational improvements
   - Close with a forward-looking statement about next steps
   - Keep under one page (150-250 words)

8. POTENTIAL STREAMLINE APPLICATIONS
   - Evaluate the client's pain points, business context, and your analysis to identify 3-5 practical workflow automations using Intellistack Streamline
   - Streamline steps: Forms, Documents, Collaboration, Sign, Notifications, Logic, Transform, Data Search, Deliver Data, Extract Data from Files, Incoming Webhook, Outbound Webhook
   - Streamline integrations: Salesforce, Google Drive, Google Sheets, Google Calendar, Amazon S3, Dropbox, OneDrive, SharePoint, Excel Online, Slack, SendGrid, SMTP, SFTP, Twilio, Microsoft Outlook
   - For each application: title, business problem (in their language), workflow steps used, integrations that apply, operational outcome
   - Rank by ease of implementation and business impact -- low-hanging fruit first
   - Use plain business language -- no technical jargon
   - NEVER include cost estimates, pricing, or timelines
   - Keep under 400 words total

OUTPUT FORMAT:
Return ONLY valid JSON in this exact structure. The "summary", "architecture_diagram", and "bottom_line" fields contain plain text. Schema "columns" use table format. All text fields can include newline characters (\\n) for formatting:
{{
  "status": "complete",
  "summary": "Executive summary text with specific data references...",
  "problems": [
    {{
      "title": "Problem Title",
      "severity": "high|medium|low",
      "evidence": "Specific evidence citing data: row counts, dollar amounts, percentages from documents...",
      "recommendation": "Concrete action: do X, expect Y outcome, costs approximately Z..."
    }}
  ],
  "architecture_diagram": "ASCII diagram showing proposed system architecture using +, -, |, v, > characters...",
  "schema": {{
    "tables": [
      {{
        "name": "table_name",
        "purpose": "what this table manages",
        "columns": [
          {{"name": "column_name", "type": "data_type", "description": "what it stores"}}
        ]
      }}
    ],
    "relationships": [
      "table_a.column -> table_b.column (one-to-many)"
    ]
  }},
  "plan": [
    {{
      "phase": "7-day: Build & Demo",
      "actions": ["1. Specific action with measurable outcome", "2. Another action"]
    }},
    {{
      "phase": "14-day: Validate & Connect",
      "actions": ["1. Specific action with measurable outcome", "2. Another action"]
    }},
    {{
      "phase": "21-day: Deploy or Decide",
      "actions": ["1. Specific action with measurable outcome", "2. Another action"]
    }}
  ],
  "bottom_line": "Direct summary: what to do first, what it costs, what to expect...",
  "client_summary": "Based on the information provided, XO has identified the following opportunities for [Company Name]:\\n\\n- First value proposition as a business outcome\\n- Second value proposition\\n- Third value proposition\\n\\nForward-looking closing statement about next steps.",
  "streamline_applications": "Based on [Company Name]'s operational needs, Streamline can automate the following workflows:\\n\\n**1. [Application Title]**\\nProblem: [Business problem in their language]\\nWorkflow: [Steps used e.g. Forms → Logic → Documents → Sign → Notifications]\\nIntegrations: [e.g. Salesforce, Google Drive, Slack]\\nOutcome: [What changes day-to-day]\\n\\n**2. [Application Title]**\\n...\\n\\nThese applications are ordered by ease of implementation.",
  "sources": [
    {{"type": "client_data", "reference": "filename or data source"}}
  ]
}}

Be specific. Use actual data from the documents. Think like a management consultant presenting to the CEO."""

    try:
        message = anthropic_client.messages.create(
            model=model,
            max_tokens=16000,
            temperature=0.7,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )

        response_text = message.content[0].text
        stop_reason = message.stop_reason
        print(f"Claude response: {len(response_text)} chars, stop_reason={stop_reason}, model={model}")

        if '```json' in response_text:
            start = response_text.find('```json') + 7
            end = response_text.find('```', start)
            response_text = response_text[start:end].strip()
        elif '```' in response_text:
            start = response_text.find('```') + 3
            end = response_text.find('```', start)
            response_text = response_text[start:end].strip()

        # Try parsing JSON, with repair for truncated responses
        try:
            analysis = json.loads(response_text)
        except json.JSONDecodeError as je:
            print(f"JSON parse failed at char {je.pos}, attempting repair...")
            analysis = _repair_truncated_json(response_text)

        analysis['analyzed_at'] = datetime.now(timezone.utc).isoformat()
        analysis['analyzed_files'] = list(extracted_text.keys())

        return analysis

    except Exception as e:
        print(f"Error calling Claude API: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "error": str(e),
            "summary": "Analysis failed",
            "problems": [],
            "schema": {"tables": []},
            "plan": [],
            "sources": []
        }


def _repair_truncated_json(text):
    """
    Attempt to repair truncated JSON from Claude (e.g. when max_tokens was hit).
    Closes unclosed strings, arrays, and objects to produce valid JSON.
    """
    # Strip trailing whitespace/incomplete tokens
    text = text.rstrip()

    # If we're inside an unclosed string, close it
    # Count unescaped quotes to see if we're mid-string
    in_string = False
    i = 0
    while i < len(text):
        c = text[i]
        if c == '\\' and in_string:
            i += 2  # skip escaped char
            continue
        if c == '"':
            in_string = not in_string
        i += 1

    if in_string:
        # Truncate back to the last complete line before the unclosed string,
        # or just close the string
        text += '"'

    # Now close any unclosed brackets/braces
    # Count open vs close for [ ] and { }
    stack = []
    in_str = False
    i = 0
    while i < len(text):
        c = text[i]
        if c == '\\' and in_str:
            i += 2
            continue
        if c == '"':
            in_str = not in_str
        elif not in_str:
            if c in ('{', '['):
                stack.append(c)
            elif c == '}' and stack and stack[-1] == '{':
                stack.pop()
            elif c == ']' and stack and stack[-1] == '[':
                stack.pop()
        i += 1

    # Remove any trailing comma before we close brackets
    text = text.rstrip()
    if text.endswith(','):
        text = text[:-1]

    # Close unclosed brackets in reverse order
    for bracket in reversed(stack):
        if bracket == '{':
            text += '}'
        elif bracket == '[':
            text += ']'

    try:
        result = json.loads(text)
        print(f"JSON repair succeeded — recovered {len(str(result))} chars of analysis")
        return result
    except json.JSONDecodeError as e2:
        print(f"JSON repair failed: {str(e2)}")
        # Last resort: try to extract whatever we can
        raise Exception(f"Could not parse Claude response even after repair: {str(e2)}")
